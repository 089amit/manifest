print("App is starting...")
from flask import Flask, request, render_template, send_file, jsonify, redirect, url_for, session
from functools import wraps
import pandas as pd
import os
import uuid
import shutil
import threading
import time
import random
import json
import re
from tracking import tracking_bp
from datetime import datetime
from openpyxl.utils import get_column_letter
from openpyxl.styles import Alignment, Font, Border, Side
import openpyxl
from num2words import num2words
from docx import Document
import zipfile
import io

# ---------- Tracking Blueprint ----------
from tracking import tracking_bp

app = Flask(__name__)
app.secret_key = os.environ.get('SECRET_KEY', 'your-secret-key-change-in-production')

UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'outputs'
TEMP_FOLDER = 'temp_files'
# Not cleaned up by cleanup_temp_files() - holds small persistent settings like
# the last-used manifest/invoice number, unlike TEMP_FOLDER which is wiped hourly.
PERSISTENT_DATA_FOLDER = 'persistent_data'

# Serverless hosts like Vercel have a read-only filesystem everywhere except
# /tmp. Rather than rely on a platform env var (which may not even be exposed
# unless explicitly enabled in project settings), detect this directly: try
# the normal relative folders, and only fall back to /tmp if that actually
# fails. /tmp there is writable but ephemeral and not shared across function
# instances - see deployment notes for why a traditional host suits this
# app's session-based workflow better.
IS_VERCEL = False
try:
    os.makedirs(UPLOAD_FOLDER, exist_ok=True)
    os.makedirs(OUTPUT_FOLDER, exist_ok=True)
    os.makedirs(TEMP_FOLDER, exist_ok=True)
    os.makedirs(PERSISTENT_DATA_FOLDER, exist_ok=True)
except OSError:
    IS_VERCEL = True
    UPLOAD_FOLDER = os.path.join('/tmp', UPLOAD_FOLDER)
    OUTPUT_FOLDER = os.path.join('/tmp', OUTPUT_FOLDER)
    TEMP_FOLDER = os.path.join('/tmp', TEMP_FOLDER)
    PERSISTENT_DATA_FOLDER = os.path.join('/tmp', PERSISTENT_DATA_FOLDER)
    os.makedirs(UPLOAD_FOLDER, exist_ok=True)
    os.makedirs(OUTPUT_FOLDER, exist_ok=True)
    os.makedirs(TEMP_FOLDER, exist_ok=True)
    os.makedirs(PERSISTENT_DATA_FOLDER, exist_ok=True)

LAST_MANIFEST_NUMBER_FILE = os.path.join(PERSISTENT_DATA_FOLDER, 'last_manifest_number.json')

def get_last_manifest_number():
    """Return the last manifest/invoice number saved, or '' if none yet."""
    if os.path.exists(LAST_MANIFEST_NUMBER_FILE):
        try:
            with open(LAST_MANIFEST_NUMBER_FILE, 'r') as f:
                return json.load(f).get('last_manifest_number', '')
        except (json.JSONDecodeError, OSError):
            return ''
    return ''

def save_last_manifest_number(value):
    """Persist the last manifest/invoice number so the user can continue from there next time."""
    if not value:
        return
    try:
        with open(LAST_MANIFEST_NUMBER_FILE, 'w') as f:
            json.dump({'last_manifest_number': str(value), 'updated_at': datetime.now().isoformat()}, f)
    except OSError:
        app.logger.exception("Failed to save last manifest number")

LAST_BOX_LIMIT_FILE = os.path.join(PERSISTENT_DATA_FOLDER, 'last_box_limit.json')
DEFAULT_MAX_BOX_PER_PART = 10  # used only if nothing has ever been saved

def get_last_box_limit():
    """Return the last-used 'boxes per custom manifest part' limit, or the
    default if it's never been set. Lets customs limits be changed from the
    UI each month instead of editing code."""
    if os.path.exists(LAST_BOX_LIMIT_FILE):
        try:
            with open(LAST_BOX_LIMIT_FILE, 'r') as f:
                return int(json.load(f).get('max_box_per_part', DEFAULT_MAX_BOX_PER_PART))
        except (json.JSONDecodeError, OSError, ValueError, TypeError):
            return DEFAULT_MAX_BOX_PER_PART
    return DEFAULT_MAX_BOX_PER_PART

def save_last_box_limit(value):
    """Persist the box-per-part limit so it's remembered next time."""
    try:
        with open(LAST_BOX_LIMIT_FILE, 'w') as f:
            json.dump({'max_box_per_part': int(value), 'updated_at': datetime.now().isoformat()}, f)
    except (OSError, ValueError, TypeError):
        app.logger.exception("Failed to save last box limit")

def get_session_bag_markings(session_dir):
    """Read the persisted {HAWB: marking} map for a session, if any."""
    path = os.path.join(session_dir, 'bag_markings.json')
    if os.path.exists(path):
        try:
            with open(path, 'r') as f:
                return json.load(f)
        except (json.JSONDecodeError, OSError):
            return {}
    return {}

def save_session_bag_markings(session_dir, mapping):
    """Persist the {HAWB: marking} map for a session (only non-empty values)."""
    path = os.path.join(session_dir, 'bag_markings.json')
    clean = {h: b for h, b in mapping.items() if b not in (None, '')}
    try:
        with open(path, 'w') as f:
            json.dump(clean, f)
    except OSError:
        app.logger.exception("Failed to save session bag markings")

download_files = {}
chamber_download_files = {}
session_merged_data = {}

# ---------- Authentication ----------
ADMIN_USERNAME = os.environ.get('ADMIN_USERNAME', 'admin')
ADMIN_PASSWORD = os.environ.get('ADMIN_PASSWORD', 'admin123')

def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not session.get('logged_in'):
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        if username == ADMIN_USERNAME and password == ADMIN_PASSWORD:
            session['logged_in'] = True
            return redirect(url_for('index'))
        else:
            return render_template('login.html', error='Invalid credentials')
    return render_template('login.html')

@app.route('/logout')
def logout():
    session.pop('logged_in', None)
    return redirect(url_for('login'))

# ---------- Helper: placeholder replacement ----------
def replace_placeholders_in_doc(doc, replacements):
    replace_count = 0
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                for run in paragraph.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, value)
                        replace_count += 1
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, value)
                    replace_count += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in replacements.items():
                        if key in paragraph.text:
                            for run in paragraph.runs:
                                if key in run.text:
                                    run.text = run.text.replace(key, value)
                                    replace_count += 1
                            if key in paragraph.text:
                                paragraph.text = paragraph.text.replace(key, value)
                                replace_count += 1
    return replace_count

# ---------- Manifest cleaning ----------
def clean_manifest(df_raw):
    df = df_raw.copy()
    header_row_idx = None
    for i, row in df.iterrows():
        first_cell = str(row.iloc[0]).strip().upper()
        if first_cell == 'S.NO':
            header_row_idx = i
            break
    if header_row_idx is None:
        raise ValueError("Could not find header row (S.NO) in manifest file")
    df.columns = df.iloc[header_row_idx]
    df = df.iloc[header_row_idx + 1:].reset_index(drop=True)
    df = df[pd.to_numeric(df['S.NO'], errors='coerce').notna()]
    needed = {
        'S.NO': 'S.NO',
        'CONSIGNER': 'CONSIGNER',
        'HAWB#': 'HAWB#',
        'GBP': 'GBP',
        'Actual WT': 'Actual WT'
    }
    col_mapping = {}
    for std_name, target in needed.items():
        for col in df.columns:
            if col.strip().upper() == std_name.upper():
                col_mapping[std_name] = col
                break
        if std_name not in col_mapping:
            raise ValueError(f"Column '{std_name}' not found in manifest file")
    result = df[[col_mapping[n] for n in needed]].copy()
    result.columns = list(needed.keys())
    result['S.NO'] = pd.to_numeric(result['S.NO'], errors='coerce').astype('Int64')
    result['Actual WT'] = pd.to_numeric(result['Actual WT'], errors='coerce')
    result['GBP'] = pd.to_numeric(result['GBP'], errors='coerce')
    return result

def load_mawb(df):
    needed = {
        'No': 'No',
        'HAWB': 'HAWB',
        'NO OF PCS': 'NO OF PCS',
        'Consignee Name': 'Consignee Name',
        'Goods Description': 'Goods Description',
        'Country': 'Destination'
    }
    col_mapping = {}
    for std_name, target in needed.items():
        for col in df.columns:
            if col.strip().upper() == std_name.upper():
                col_mapping[std_name] = col
                break
        if std_name not in col_mapping:
            raise ValueError(f"Column '{std_name}' not found in MAWB file")
    result = df[[col_mapping[n] for n in needed]].copy()
    result.columns = list(needed.keys())
    result.rename(columns={'Country': 'Destination'}, inplace=True)
    result['NO OF PCS'] = pd.to_numeric(result['NO OF PCS'], errors='coerce').fillna(1).astype(int)
    return result

def is_bag_marking(value):
    """True for any non-empty bag/box marking value from the dedicated
    'BagMarking' field (e.g. '1', '5-9')."""
    if value is None:
        return False
    return str(value).strip() != ''

def extract_hawb_to_bag(edited_rows, session_dir=None):
    """
    Build {HAWB: marking} from the WT-box report's dedicated 'BagMarking'
    field (never from 'Box', which is pcs count and must never be touched).
    Falls back to the persisted per-session marking file for any row that
    doesn't include the field, so older clients / partial payloads still work.
    """
    persisted = get_session_bag_markings(session_dir) if session_dir else {}
    hawb_to_bag = {}
    for row in edited_rows:
        hawb = row.get('HAWB')
        bag = row.get('BagMarking', row.get('Bag Marking'))
        if bag is None or str(bag).strip() == '':
            bag = persisted.get(hawb)
        bag = str(bag).strip() if bag not in (None, '') else None
        hawb_to_bag[hawb] = bag
    return hawb_to_bag

def is_valuable_item(description):
    valuable_keywords = [
        'gold', 'silver', 'statue', 'jewellery', 'jewelry', 'metal', 'brass', 'copper',
        'decoration', 'decorative', 'singing bowl', 'buddha', 'tara', 'crystal',
        'stone', 'khukuri', 'knife', 'painting', 'wooden mask', 'handicraft'
    ]
    desc_lower = description.lower()
    return any(keyword in desc_lower for keyword in valuable_keywords)

# ---------- Content-type categorisation (for customs-friendly grouping) ----------
# Order matters: checked top-to-bottom, first match wins. More specific keyword
# sets are checked before generic ones (Handicraft has broad words like
# 'metal'/'stone'/'wood', so it is checked last, as a catch-all before 'Others').
CATEGORY_DEFINITIONS = [
    ('Artificial Jewellery', [
        'jewellery', 'jewelry', 'necklace', 'earring', 'earrings', 'bracelet',
        'bangle', 'bangles', 'bead', 'beads', 'pendant', 'anklet',
        'ornament', 'ornaments', 'gem'
    ]),
    ('Cotton Dress', [
        'cotton', 'dress', 'shirt', 'trouser', 'trousers', 'pant', 'pants', 'kurta',
        'kurti', 'scarf', 'scarves', 'shawl', 'pashmina', 'wool', 'woolen', 'woollen',
        'garment', 'garments', 'clothing', 'clothes', 'apparel', 'tunic', 'skirt',
        'blouse', 'jacket', 'sweater', 'fabric', 'textile'
    ]),
    ('Food Items', [
        'tea', 'coffee', 'spice', 'spices', 'masala', 'food', 'snack', 'snacks',
        'noodle', 'noodles', 'pickle', 'honey', 'herb', 'herbs', 'seed', 'seeds',
        'grain', 'grains', 'chocolate', 'rice', 'dal', 'lentil', 'dry fruit'
    ]),
    ('Accessories', [
        'accessory', 'accessories', 'belt', 'wallet', 'purse', 'watch',
        'sunglasses', 'cap', 'hat', 'glove', 'shoe'
    ]),
    ('Handicraft', [
        'handicraft', 'statue', 'wooden', 'wood', 'mask', 'singing bowl', 'buddha',
        'tara', 'brass', 'copper', 'metal', 'carving', 'sculpture', 'decoration',
        'decorative', 'painting', 'khukuri', 'knife', 'crystal', 'stone', 'thanka',
        'bamboo', 'doko', 'pottery', 'ceramic', 'incense', 'gold', 'silver'
    ]),
]
CATEGORY_ORDER = [name for name, _ in CATEGORY_DEFINITIONS] + ['Others']

def categorize_description(description):
    """Classify a Goods Description string into a customs-friendly content category."""
    if pd.isna(description):
        return 'Others'
    desc_lower = str(description).strip().lower()
    if desc_lower == '':
        return 'Others'
    for category, keywords in CATEGORY_DEFINITIONS:
        if any(keyword in desc_lower for keyword in keywords):
            return category
    return 'Others'

def create_itemised_sheet(merged_df, use_bag_labels=False, bag_map=None, console_bag_notes=None):
    """
    console_bag_notes: optional dict {HAWB: bag_label}. When a HAWB has an
    assigned console bag number, an extra note row ("Console Bag: <label>")
    is inserted directly under that HAWB's item rows in the packing list.
    This is independent of use_bag_labels/bag_map, which fully collapses
    rows into bag-level entries for the bag-based invoice mode.
    """
    rows = []
    merged_sorted = merged_df.sort_values('HAWB#').reset_index(drop=True)
    for idx, row in merged_sorted.iterrows():
        hawb = row['HAWB#']
        desc = row['Goods Description']
        total_usd = float(row['GBP']) if pd.notna(row['GBP']) else 0.0
        if pd.isna(desc) or desc == '':
            continue
        items = [item.strip() for item in str(desc).split(',') if item.strip()]
        if not items:
            continue
        n_items = len(items)
        if n_items == 0 or total_usd == 0:
            for item in items:
                display_hawb = bag_map.get(hawb, hawb) if use_bag_labels and bag_map else hawb
                rows.append({
                    'HAWB': display_hawb if item == items[0] else '',
                    'DESCRIPTION': item,
                    'PCS': '',
                    'UNIT VALUE': '',
                    'TOTAL': '',
                    'HSCODE': ''
                })
            bag_label = console_bag_notes.get(hawb) if console_bag_notes else None
            if bag_label:
                rows.append({
                    'HAWB': '', 'DESCRIPTION': f'Console Bag: {bag_label}',
                    'PCS': '', 'UNIT VALUE': '', 'TOTAL': '', 'HSCODE': ''
                })
            if idx != len(merged_sorted) - 1:
                rows.append({'HAWB': '', 'DESCRIPTION': '', 'PCS': '', 'UNIT VALUE': '', 'TOTAL': '', 'HSCODE': ''})
            continue
        
        pcs_list = [1 if is_valuable_item(item) else random.randint(2, 15) for item in items]
        raw_unit_values = [random.uniform(1, 100) for _ in range(n_items)]
        raw_total = sum(raw_unit_values[i] * pcs_list[i] for i in range(n_items))
        scale = total_usd / raw_total if raw_total != 0 else 1
        final_unit_values = [round(raw * scale, 2) for raw in raw_unit_values]
        final_totals = [final_unit_values[i] * pcs_list[i] for i in range(n_items)]
        sum_totals = sum(final_totals)
        if abs(sum_totals - total_usd) > 0.01:
            diff = total_usd - sum_totals
            last_index = n_items - 1
            new_last_total = final_totals[last_index] + diff
            if pcs_list[last_index] != 0:
                new_last_unit = new_last_total / pcs_list[last_index]
                final_unit_values[last_index] = round(new_last_unit, 2)
                final_totals[last_index] = round(new_last_unit * pcs_list[last_index], 2)
        
        for i, item in enumerate(items):
            display_hawb = bag_map.get(hawb, hawb) if use_bag_labels and bag_map else hawb
            rows.append({
                'HAWB': display_hawb if i == 0 else '',
                'DESCRIPTION': item,
                'PCS': pcs_list[i],
                'UNIT VALUE': final_unit_values[i],
                'TOTAL': final_totals[i],
                'HSCODE': ''
            })
        bag_label = console_bag_notes.get(hawb) if console_bag_notes else None
        if bag_label:
            rows.append({
                'HAWB': '', 'DESCRIPTION': f'Console Bag: {bag_label}',
                'PCS': '', 'UNIT VALUE': '', 'TOTAL': '', 'HSCODE': ''
            })
        if idx != len(merged_sorted) - 1:
            rows.append({'HAWB': '', 'DESCRIPTION': '', 'PCS': '', 'UNIT VALUE': '', 'TOTAL': '', 'HSCODE': ''})
    return pd.DataFrame(rows)

DEFAULT_MAX_BAG_WEIGHT = 30.0  # kg - matches "a bag tops out around 30kg" consolidation rule

def suggest_bag_markings(merged_df, max_bag_weight=DEFAULT_MAX_BAG_WEIGHT):
    """
    Auto-suggest a bag/box marking per HAWB (editable/overridable by the
    user afterwards), following the real-world courier rule:
      - Small HAWBs (single box, weight well under the cap) are consolidated
        together into one bag, filling it up to ~max_bag_weight (kg); every
        HAWB in that group shares the SAME single marking number.
      - A HAWB with multiple boxes of its own (NO OF PCS > 1), or whose own
        weight already meets/exceeds max_bag_weight, is NOT merged with
        others. Instead it gets its own sequential marking: a single number
        if it's one box, or a number RANGE if it has several boxes (e.g. a
        5-box HAWB continuing from marking 5 becomes "5-9").
      - One running counter is shared across the whole manifest, so bag
        numbers and box-range numbers interleave in a single sequence.
    Returns {HAWB: marking_string}, e.g. {'H1': '1', 'H2': '1', 'H3': '5-9'}.
    """
    suggestions = {}
    counter = 1
    group_hawbs = []
    group_weight = 0.0

    def flush_group():
        nonlocal counter, group_hawbs, group_weight
        if not group_hawbs:
            return
        label = str(counter)
        for h in group_hawbs:
            suggestions[h] = label
        counter += 1
        group_hawbs = []
        group_weight = 0.0

    for _, row in merged_df.sort_values('HAWB#').iterrows():
        hawb = row['HAWB#']
        weight = float(row['Actual WT']) if pd.notna(row['Actual WT']) else 0.0
        pcs = int(row['NO OF PCS']) if pd.notna(row['NO OF PCS']) else 1

        if pcs > 1 or weight >= max_bag_weight:
            flush_group()
            start = counter
            end = counter + pcs - 1
            suggestions[hawb] = str(start) if pcs == 1 else f"{start}-{end}"
            counter = end + 1
            continue

        if group_hawbs and group_weight + weight > max_bag_weight:
            flush_group()

        group_hawbs.append(hawb)
        group_weight += weight

    flush_group()
    return suggestions

def prepare_wtbox_data(merged_df):
    bag_suggestions = suggest_bag_markings(merged_df)
    wt_data = []
    for _, row in merged_df.iterrows():
        hawb = row['HAWB#']
        sender = row['CONSIGNER']
        receiver = row['Consignee Name']
        details = f"{sender} → {receiver}"
        original_pcs = int(row['NO OF PCS'])
        weight = float(row['Actual WT'])
        country = row['Destination']
        wt_data.append({
            'HAWB': hawb,
            'Details': details,
            'Box': original_pcs,
            'BagMarking': bag_suggestions.get(hawb, ''),
            'Weight': weight,
            'Country': country
        })
    df_wt = pd.DataFrame(wt_data)
    country_counts = df_wt['Country'].value_counts().to_dict()
    df_wt['_count'] = df_wt['Country'].map(country_counts)
    df_wt = df_wt.sort_values('_count', ascending=True).drop('_count', axis=1)
    return df_wt

MAX_USD_PER_PART = 950
MAX_WEIGHT_PER_PART = 1500

def split_manifest(merged_df, bag_map=None, max_box_per_part=None):
    """
    Build manifest parts with two priorities, in this order:
      1) TOP PRIORITY: each part should contain exactly max_box_per_part
         boxes wherever possible, even if that means mixing content
         categories together.
      2) SECOND PRIORITY: prefer keeping the same content category together
         within a part, for easier customs inspection.
    USD (950) and weight (1500 kg) remain hard caps per part.

    max_box_per_part: boxes-per-part limit for this run. Passed in by the
    caller (sourced from the request / last-saved setting) so this can be
    changed from the UI when customs limits change, instead of editing
    code. Falls back to DEFAULT_MAX_BOX_PER_PART if not given.

    bag_map: optional {HAWB: bag_label} dict. When provided, the packing
    list for each part will show a "Console Bag: <label>" note directly
    under the relevant HAWB's line items.
    """
    max_box_per_part = int(max_box_per_part) if max_box_per_part else DEFAULT_MAX_BOX_PER_PART
    df = merged_df.reset_index(drop=True).copy()
    df['_category'] = df['Goods Description'].apply(categorize_description)

    category_queues = {cat: list(df.index[df['_category'] == cat]) for cat in CATEGORY_ORDER}

    def any_left():
        return any(category_queues[c] for c in CATEGORY_ORDER)

    parts = []
    itemised_parts = []

    while any_left():
        bin_indices = []
        cum_pcs = 0
        cum_usd = 0.0
        cum_weight = 0.0

        def try_add(idx):
            nonlocal cum_pcs, cum_usd, cum_weight
            row = df.loc[idx]
            pcs = int(row['NO OF PCS'])
            usd = float(row['GBP'])
            weight = float(row['Actual WT'])
            if bin_indices and (cum_pcs + pcs > max_box_per_part or cum_usd + usd > MAX_USD_PER_PART or cum_weight + weight > MAX_WEIGHT_PER_PART):
                return False
            bin_indices.append(idx)
            cum_pcs += pcs
            cum_usd += usd
            cum_weight += weight
            return True

        primary = next((c for c in CATEGORY_ORDER if category_queues[c]), None)

        while category_queues[primary] and cum_pcs < max_box_per_part:
            idx = category_queues[primary][0]
            if try_add(idx):
                category_queues[primary].pop(0)
            else:
                break

        if cum_pcs < max_box_per_part:
            for cat in CATEGORY_ORDER:
                if cat == primary:
                    continue
                while category_queues[cat] and cum_pcs < max_box_per_part:
                    idx = category_queues[cat][0]
                    if try_add(idx):
                        category_queues[cat].pop(0)
                    else:
                        break

        part_df = df.loc[bin_indices].drop(columns=['_category']).reset_index(drop=True)
        parts.append(part_df)
        itemised_parts.append(create_itemised_sheet(part_df, console_bag_notes=bag_map))

    return parts, itemised_parts

def save_excel_for_part(part_df, itemised_df, part_number, invoice_date, temp_dir, invoice_number, narration_text=None):
    wb = openpyxl.Workbook()
    default_sheet = wb.active
    wb.remove(default_sheet)
    
    box_start = 1
    box_ranges = []
    total_boxes = 0
    total_weight = 0.0
    for _, row in part_df.iterrows():
        pcs = int(row['NO OF PCS'])
        total_boxes += pcs
        weight = float(row['Actual WT'])
        total_weight += weight
        end = box_start + pcs - 1
        box_ranges.append(f"{box_start}-{end}" if pcs > 1 else str(box_start))
        box_start = end + 1
    part_df['BOX_RANGE'] = box_ranges
    
    manifest_data = part_df[[
        'BOX_RANGE', 'HAWB#', 'NO OF PCS', 'Actual WT', 'GBP', 'CONSIGNER', 'Consignee Name', 'Goods Description', 'Destination'
    ]].copy()
    manifest_data.columns = ['BOX', 'HAWB', 'PCS', 'WEIGHT', 'USD', 'SENDER', 'RECEIVER', 'DESCRIPTION', 'DESTINATION']
    manifest_data['REMARKS'] = ''
    manifest_data['PCS'] = pd.to_numeric(manifest_data['PCS'], errors='coerce').fillna(0).astype(int)
    manifest_data['WEIGHT'] = pd.to_numeric(manifest_data['WEIGHT'], errors='coerce').fillna(0.0)
    manifest_data['USD'] = pd.to_numeric(manifest_data['USD'], errors='coerce').fillna(0.0)
    
    total_row = {'BOX': 'TOTAL', 'PCS': manifest_data['PCS'].sum(), 'WEIGHT': manifest_data['WEIGHT'].sum(), 'USD': manifest_data['USD'].sum()}
    for col in manifest_data.columns:
        if col not in total_row:
            total_row[col] = ''
    manifest_with_total = pd.concat([manifest_data, pd.DataFrame([total_row])], ignore_index=True)
    
    ws_manifest = wb.create_sheet(f'Custom Manifest')
    ws_manifest.page_setup.orientation = 'landscape'
    ws_manifest.page_setup.paperSize = 9
    ws_manifest.page_setup.fitToPage = True
    ws_manifest.page_setup.fitToWidth = 1
    ws_manifest.page_setup.fitToHeight = 0
    
    max_col = len(manifest_with_total.columns)
    start_row = add_title_block_sheet1(ws_manifest, total_boxes, total_weight, invoice_date, max_col, invoice_number)
    last_row = write_data_to_sheet(ws_manifest, manifest_with_total, start_row)
    apply_borders(ws_manifest, last_row, max_col)
    auto_fit_columns(ws_manifest, start_row)
    
    if narration_text:
        narration_row = last_row + 2
        ws_manifest.merge_cells(start_row=narration_row, start_column=1, end_row=narration_row, end_column=max_col)
        cell = ws_manifest.cell(row=narration_row, column=1, value=narration_text)
        cell.alignment = Alignment(wrap_text=True, vertical='top', horizontal='left')
        cell.font = Font(italic=True, size=10)
        line_count = narration_text.count('\n') + 1
        ws_manifest.row_dimensions[narration_row].height = 15 * line_count
        last_row = narration_row
    
    if not itemised_df.empty and 'TOTAL' in itemised_df.columns:
        total_row_packing = {}
        for col in itemised_df.columns:
            if col == 'PCS':
                total_val = pd.to_numeric(itemised_df[col], errors='coerce').sum()
                total_row_packing[col] = int(total_val) if not pd.isna(total_val) else 0
            elif col == 'TOTAL':
                total_val = pd.to_numeric(itemised_df[col], errors='coerce').sum()
                total_row_packing[col] = round(total_val, 2)
            else:
                total_row_packing[col] = ''
        total_row_packing['DESCRIPTION'] = 'TOTAL'
        total_row_packing['HAWB'] = ''
        packing_with_total = pd.concat([itemised_df, pd.DataFrame([total_row_packing])], ignore_index=True)
    else:
        packing_with_total = itemised_df.copy()
    
    ws_packing = wb.create_sheet(f'Custom Packing List')
    ws_packing.page_setup.orientation = 'portrait'
    ws_packing.page_setup.paperSize = 9
    ws_packing.page_setup.fitToPage = True
    ws_packing.page_setup.fitToWidth = 1
    ws_packing.page_setup.fitToHeight = 0
    
    max_col_p = len(packing_with_total.columns)
    start_row_p = add_title_block_sheet2(ws_packing, total_boxes, total_weight, invoice_date, max_col_p, invoice_number)
    last_row_p = write_data_to_sheet(ws_packing, packing_with_total, start_row_p)
    apply_borders(ws_packing, last_row_p, max_col_p)
    auto_fit_columns(ws_packing, start_row_p)
    
    safe_inv = re.sub(r'[\\/*?:"<>|]', '_', invoice_number)
    filename = f"manifest-{safe_inv}.xlsx"
    filepath = os.path.join(temp_dir, filename)
    wb.save(filepath)
    os.chmod(filepath, 0o644)
    return filepath, filename

def add_title_block_sheet1(worksheet, total_boxes, total_weight, invoice_date, max_col, invoice_number):
    worksheet.insert_rows(1, amount=13)
    worksheet.merge_cells(start_row=1, start_column=1, end_row=1, end_column=max_col)
    worksheet['A1'] = 'DOKO EXPORTS PVT LTD'
    worksheet['A1'].font = Font(bold=True, size=14)
    worksheet['A1'].alignment = Alignment(horizontal='center', vertical='center')
    worksheet.merge_cells(start_row=2, start_column=1, end_row=2, end_column=max_col)
    worksheet['A2'] = 'INVOICE/PACKING LIST'
    worksheet['A2'].font = Font(bold=True, size=12)
    worksheet['A2'].alignment = Alignment(horizontal='center', vertical='center')
    worksheet['A5'] = f'Invoice No:- {invoice_number}'
    worksheet['A5'].font = Font(bold=True)
    worksheet['G5'] = 'EXIM NO:- 6197818390116NP'
    worksheet['G5'].font = Font(bold=True)
    worksheet['A6'] = 'Consignee Name, Address & Country'
    worksheet['A6'].font = Font(bold=True)
    worksheet['G6'] = f'Invoice Date :- {invoice_date}'
    worksheet['G6'].font = Font(bold=True)
    worksheet['G7'] = 'PAN NO: 619781839'
    worksheet['G7'].font = Font(bold=True)
    worksheet['A8'] = 'To:'
    worksheet['A8'].font = Font(bold=True)
    worksheet['B8'] = 'ESHIPPER EXPRESS COURIER'
    worksheet['G8'] = f'Box- {total_boxes} boxes'
    worksheet['G8'].font = Font(bold=True)
    worksheet['B9'] = 'UK'
    worksheet['G9'] = f'Weight: {total_weight:.2f} kgs'
    worksheet['G9'].font = Font(bold=True)
    worksheet['G10'] = 'Shipment by : Air Freight'
    worksheet['G10'].font = Font(bold=True)
    return 14

def add_title_block_sheet2(worksheet, total_boxes, total_weight, invoice_date, max_col, invoice_number):
    worksheet.insert_rows(1, amount=13)
    worksheet.merge_cells(start_row=1, start_column=1, end_row=1, end_column=max_col)
    worksheet['A1'] = 'DOKO EXPORTS PVT LTD'
    worksheet['A1'].font = Font(bold=True, size=14)
    worksheet['A1'].alignment = Alignment(horizontal='center', vertical='center')
    worksheet.merge_cells(start_row=2, start_column=1, end_row=2, end_column=max_col)
    worksheet['A2'] = 'INVOICE/PACKING LIST'
    worksheet['A2'].font = Font(bold=True, size=12)
    worksheet['A2'].alignment = Alignment(horizontal='center', vertical='center')
    worksheet['A5'] = f'Invoice No:- {invoice_number}'
    worksheet['A5'].font = Font(bold=True)
    worksheet['C5'] = f'Invoice Date :- {invoice_date}'
    worksheet['C5'].font = Font(bold=True)
    worksheet['A6'] = 'Consignee Name, Address & Country'
    worksheet['A6'].font = Font(bold=True)
    worksheet['C6'] = 'EXIM NO:-6197818390116NP'
    worksheet['C6'].font = Font(bold=True)
    worksheet['C7'] = 'PAN NO:619781839'
    worksheet['C7'].font = Font(bold=True)
    worksheet['A8'] = 'To:'
    worksheet['A8'].font = Font(bold=True)
    worksheet['B8'] = 'ESHIPPER EXPRESS COURIER'
    worksheet['C8'] = f'Box- {total_boxes} boxes'
    worksheet['C8'].font = Font(bold=True)
    worksheet['B9'] = 'UK'
    worksheet['C9'] = f'Weight: {total_weight:.2f} kgs'
    worksheet['C9'].font = Font(bold=True)
    worksheet['C10'] = 'Shipment by : Air Freight'
    worksheet['C10'].font = Font(bold=True)
    return 14

def apply_borders(worksheet, max_row, max_col):
    thin_border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))
    for row in range(1, max_row + 1):
        for col in range(1, max_col + 1):
            worksheet.cell(row=row, column=col).border = thin_border

def auto_fit_columns(worksheet, start_row):
    for column in worksheet.columns:
        max_length = 0
        col_letter = get_column_letter(column[0].column)
        for cell in column[start_row-1:]:
            try:
                if cell.value:
                    if isinstance(cell.value, str) and '\n' in cell.value:
                        lines = cell.value.split('\n')
                        cell_length = max(len(line) for line in lines)
                    else:
                        cell_length = len(str(cell.value))
                    if cell_length > max_length:
                        max_length = cell_length
            except:
                pass
        adjusted_width = min(max_length + 2, 50)
        worksheet.column_dimensions[col_letter].width = adjusted_width
        for cell in column[start_row-1:]:
            cell.alignment = Alignment(wrap_text=True, vertical='top')

def write_data_to_sheet(worksheet, data_df, start_row):
    header_row = start_row
    for col_idx, col_name in enumerate(data_df.columns, 1):
        cell = worksheet.cell(row=header_row, column=col_idx, value=col_name)
        cell.font = Font(bold=True)
    for r_idx, row in enumerate(data_df.values, start=header_row+1):
        for c_idx, value in enumerate(row, 1):
            worksheet.cell(row=r_idx, column=c_idx, value=value)
    return header_row + len(data_df)

def generate_chamber_certificate(session_dir, meta, invoice_number, invoice_date, part_number):
    template_path = os.path.join(os.path.dirname(__file__), 'templates', 'CHAMBER_template.docx')
    if not os.path.exists(template_path):
        return None, None
    doc = Document(template_path)
    consignee_details = f"{meta['consignee']}, {meta['destination']}"
    total_usd_words = num2words(meta['total_usd'], lang='en', to='currency', currency='USD').upper()
    replacements = {
        '{{MANIFEST_TITLE}}': 'DOKO EXPORTS PVT LTD',
        '{{CONSIGNEE_DETAILS}}': consignee_details,
        '{{DATE}}': invoice_date,
        '{{TOTAL_USD}}': f"{meta['total_usd']:.2f}",
        '{{TOTAL_USD_WORDS}}': total_usd_words,
        '{{TOTAL_BOXES}}': str(meta['total_boxes']),
        '{{PACKING_TOTAL_PCS}}': str(meta['packing_total_pcs']),
        '{{INVOICE_NUMBER}}': invoice_number
    }
    replace_count = replace_placeholders_in_doc(doc, replacements)
    if replace_count == 0:
        return None, None
    safe_inv = re.sub(r'[\\/*?:"<>|]', '_', invoice_number)
    filename = f"chamber-{safe_inv}.docx"
    filepath = os.path.join(session_dir, filename)
    doc.save(filepath)
    os.chmod(filepath, 0o644)
    return filepath, filename

# ---------- Routes (protected) ----------
@app.route('/')
@login_required
def index():
    return render_template('index.html')

@app.route('/upload_and_prepare', methods=['POST'])
@login_required
def upload_and_prepare():
    manifest_file = request.files.get('primary')
    mawb_file = request.files.get('secondary')
    if not manifest_file or not mawb_file:
        return jsonify({'error': 'Please upload both files.'}), 400

    manifest_path = os.path.join(UPLOAD_FOLDER, str(uuid.uuid4()) + "_" + manifest_file.filename)
    mawb_path = os.path.join(UPLOAD_FOLDER, str(uuid.uuid4()) + "_" + mawb_file.filename)
    manifest_file.save(manifest_path)
    mawb_file.save(mawb_path)

    try:
        df_manifest_raw = pd.read_excel(manifest_path, header=None)
        df_mawb_raw = pd.read_excel(mawb_path)
        df_manifest = clean_manifest(df_manifest_raw)
        df_mawb = load_mawb(df_mawb_raw)
        merged = pd.merge(df_manifest, df_mawb, left_on='HAWB#', right_on='HAWB', how='inner')
        merged['Actual WT'] = pd.to_numeric(merged['Actual WT'], errors='coerce')
        merged['GBP'] = pd.to_numeric(merged['GBP'], errors='coerce')
        merged['NO OF PCS'] = pd.to_numeric(merged['NO OF PCS'], errors='coerce').fillna(1).astype(int)
        merged = merged.sort_values('HAWB#').reset_index(drop=True)

        session_id = str(uuid.uuid4())
        session_dir = os.path.join(TEMP_FOLDER, session_id)
        os.makedirs(session_dir, exist_ok=True)
        merged.to_csv(os.path.join(session_dir, 'merged.csv'), index=False)
        session_merged_data[session_id] = merged

        invoice_date = datetime.now().strftime('%d-%m-%Y')
        with open(os.path.join(session_dir, 'invoice_date.txt'), 'w') as f:
            f.write(invoice_date)

        wt_df = prepare_wtbox_data(merged)
        wt_data_json = wt_df.to_dict(orient='records')

        return jsonify({
            'session_id': session_id,
            'wtbox_data': wt_data_json
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        for p in [manifest_path, mawb_path]:
            if os.path.exists(p):
                os.remove(p)

def build_bag_groups(merged_df, hawb_to_bag):
    """Group rows of merged_df by their assigned marking. Returns
    {marking: [row, ...]} and the list of rows with no marking assigned."""
    bag_groups = {}
    unassigned_hawbs = []
    for _, row in merged_df.iterrows():
        hawb = row['HAWB#']
        bag = hawb_to_bag.get(hawb)
        if bag:
            bag_groups.setdefault(bag, []).append(row)
        else:
            unassigned_hawbs.append(row)
    return bag_groups, unassigned_hawbs

def build_bag_rows(bag_groups, unassigned_hawbs):
    """
    Turn grouped rows into the flat row-list used for manifest splitting:
      - A marking shared by MORE THAN ONE HAWB is a true consolidated bag:
        collapse to a single row (NO OF PCS=1, summed weight/usd).
      - A marking held by exactly ONE HAWB (e.g. a standalone multi-box
        HAWB tagged with a sequential range like '5-9') is NOT collapsed -
        its real pcs/weight/usd are preserved, only tagged with the marking.
      - HAWBs with no marking pass through unchanged.
    """
    bag_rows = []
    for bag, rows_list in bag_groups.items():
        if len(rows_list) > 1:
            total_weight = sum(float(r['Actual WT']) for r in rows_list)
            total_usd = sum(float(r['GBP']) for r in rows_list)
            total_pcs = sum(int(r['NO OF PCS']) for r in rows_list)
            first_row = rows_list[0]
            descriptions = set(str(r['Goods Description']) for r in rows_list if pd.notna(r['Goods Description']))
            combined_desc = ', '.join(descriptions) if descriptions else ''
            bag_rows.append({
                'BOX_RANGE': bag,
                'HAWB#': bag,
                'NO OF PCS': total_pcs,
                'Actual WT': total_weight,
                'GBP': total_usd,
                'CONSIGNER': first_row['CONSIGNER'],
                'Consignee Name': first_row['Consignee Name'],
                'Goods Description': combined_desc,
                'Destination': ''
            })
        else:
            row = rows_list[0]
            bag_rows.append({
                'BOX_RANGE': bag,
                'HAWB#': row['HAWB#'],
                'NO OF PCS': row['NO OF PCS'],
                'Actual WT': row['Actual WT'],
                'GBP': row['GBP'],
                'CONSIGNER': row['CONSIGNER'],
                'Consignee Name': row['Consignee Name'],
                'Goods Description': row['Goods Description'],
                'Destination': row['Destination']
            })
    for row in unassigned_hawbs:
        bag_rows.append({
            'BOX_RANGE': '',
            'HAWB#': row['HAWB#'],
            'NO OF PCS': row['NO OF PCS'],
            'Actual WT': row['Actual WT'],
            'GBP': row['GBP'],
            'CONSIGNER': row['CONSIGNER'],
            'Consignee Name': row['Consignee Name'],
            'Goods Description': row['Goods Description'],
            'Destination': row['Destination']
        })
    return bag_rows

@app.route('/preview_bag_parts', methods=['POST'])
@login_required
def preview_bag_parts():
    data = request.get_json()
    session_id = data.get('session_id')
    edited_rows = data.get('data')
    is_bag_based = data.get('is_bag_based', False)
    if not session_id or not edited_rows:
        return jsonify({'error': 'Invalid request'}), 400

    session_dir = os.path.join(TEMP_FOLDER, session_id)
    merged = session_merged_data.get(session_id)
    if merged is None:
        merged_path = os.path.join(session_dir, 'merged.csv')
        if not os.path.exists(merged_path):
            return jsonify({'error': 'Original data not found'}), 400
        merged = pd.read_csv(merged_path)

    hawb_to_bag = extract_hawb_to_bag(edited_rows, session_dir)

    max_box_per_part = data.get('max_box_per_part')
    if max_box_per_part in (None, ''):
        max_box_per_part = get_last_box_limit()
    else:
        try:
            max_box_per_part = int(max_box_per_part)
            if max_box_per_part <= 0:
                raise ValueError
        except (TypeError, ValueError):
            return jsonify({'error': 'max_box_per_part must be a positive integer'}), 400

    # Mirror generate_manifest's branching exactly, so the part count shown
    # here (and the number of invoice-number fields the user is asked to
    # fill in) always matches what will actually be generated - regardless
    # of whether any bag marking has been assigned.
    if not is_bag_based:
        parts, _ = split_manifest(merged, bag_map=hawb_to_bag, max_box_per_part=max_box_per_part)
        return jsonify({'part_count': len(parts)})

    bag_groups, unassigned_hawbs = build_bag_groups(merged, hawb_to_bag)
    bag_rows = build_bag_rows(bag_groups, unassigned_hawbs)

    if not bag_rows:
        return jsonify({'part_count': 0}), 200
    bag_df = pd.DataFrame(bag_rows)
    parts, _ = split_manifest(bag_df, max_box_per_part=max_box_per_part)
    return jsonify({'part_count': len(parts)})

@app.route('/generate_manifest', methods=['POST'])
@login_required
def generate_manifest():
    data = request.get_json()
    session_id = data.get('session_id')
    edited_rows = data.get('data')
    is_bag_based = data.get('is_bag_based', False)
    invoice_numbers = data.get('invoice_numbers', None)

    if not session_id or not edited_rows:
        return jsonify({'error': 'Invalid request'}), 400

    max_box_per_part = data.get('max_box_per_part')
    if max_box_per_part in (None, ''):
        max_box_per_part = get_last_box_limit()
    else:
        try:
            max_box_per_part = int(max_box_per_part)
            if max_box_per_part <= 0:
                raise ValueError
        except (TypeError, ValueError):
            return jsonify({'error': 'max_box_per_part must be a positive integer'}), 400
    save_last_box_limit(max_box_per_part)

    session_dir = os.path.join(TEMP_FOLDER, session_id)
    merged = session_merged_data.get(session_id)
    if merged is None:
        merged_path = os.path.join(session_dir, 'merged.csv')
        if not os.path.exists(merged_path):
            return jsonify({'error': 'Original data not found. Please re-upload.'}), 400
        merged = pd.read_csv(merged_path)

    hawb_to_bag = extract_hawb_to_bag(edited_rows, session_dir)
    save_session_bag_markings(session_dir, hawb_to_bag)

    if is_bag_based:
        bag_groups, unassigned_hawbs = build_bag_groups(merged, hawb_to_bag)
        bag_rows = build_bag_rows(bag_groups, unassigned_hawbs)

        if not bag_rows:
            return jsonify({'error': 'No data to generate manifest.'}), 400

        bag_df = pd.DataFrame(bag_rows)
        box_start = 1
        box_ranges = []
        for _, row in bag_df.iterrows():
            pcs = int(row['NO OF PCS'])
            end = box_start + pcs - 1
            box_ranges.append(f"{box_start}-{end}" if pcs > 1 else str(box_start))
            box_start = end + 1
        bag_df['BOX_RANGE'] = box_ranges

        parts, itemised_parts = split_manifest(bag_df, max_box_per_part=max_box_per_part)

        if invoice_numbers is None or len(invoice_numbers) != len(parts):
            app.logger.error(f"Invoice count mismatch: expected {len(parts)}, got {len(invoice_numbers) if invoice_numbers else 0}")
            return jsonify({'error': f'Invoice number count mismatch: expected {len(parts)} but received {len(invoice_numbers) if invoice_numbers else 0}. Please re-enter.'}), 400

        invoice_date_path = os.path.join(session_dir, 'invoice_date.txt')
        invoice_date = open(invoice_date_path).read().strip() if os.path.exists(invoice_date_path) else datetime.now().strftime('%d-%m-%Y')

        hawb_to_bag_for_packing = {}
        for bag, rows_list in bag_groups.items():
            for row in rows_list:
                hawb_to_bag_for_packing[row['HAWB#']] = bag

        # A marking backed by more than one HAWB is a true consolidated bag
        # (its row's HAWB# was replaced by the bag label itself); a marking
        # held by a single HAWB (e.g. a standalone multi-box range like
        # '5-9') keeps its real HAWB# and isn't a synthetic bag row.
        consolidated_labels = {bag for bag, rows_list in bag_groups.items() if len(rows_list) > 1}

        file_list = []
        heading_consignee = "ESHIPPER EXPRESS COURIER"
        heading_destination = "UK"

        for idx, (part_df, itemised_df) in enumerate(zip(parts, itemised_parts), start=1):
            inv_num = invoice_numbers[idx-1] if invoice_numbers and idx-1 < len(invoice_numbers) else f"BAG-{idx}"
            if not inv_num or inv_num.strip() == '':
                inv_num = f"BAG-{idx}"
            part_bag_narration = []
            for _, row in part_df.iterrows():
                if row['HAWB#'] in consolidated_labels:
                    bag_label = row['HAWB#']
                    hawb_list = [r['HAWB#'] for r in bag_groups.get(bag_label, [])]
                    if hawb_list:
                        part_bag_narration.append(f"{bag_label} ({', '.join(hawb_list)})")
            narration_text = '\n'.join(part_bag_narration) if part_bag_narration else ''

            # Rebuild the REAL per-HAWB rows for this part's packing list.
            # part_df has one row per bag (collapsed, deduplicated description,
            # no per-HAWB detail) - expand consolidated bags back to their
            # original HAWB rows so the itemised sheet shows genuine per-HAWB
            # items/pcs with a "Console Bag: <label>" note, exactly like the
            # non-bag-based branch does. Standalone marked HAWBs (single-row
            # "bags", e.g. box ranges) and unassigned HAWBs already ARE real
            # rows in part_df, so they pass through unchanged.
            real_rows_for_part = []
            for _, row in part_df.iterrows():
                label = row['HAWB#']
                if label in consolidated_labels:
                    real_rows_for_part.extend(bag_groups[label])
                else:
                    real_rows_for_part.append(row)
            real_part_df = pd.DataFrame(real_rows_for_part)

            itemised_df = create_itemised_sheet(real_part_df, console_bag_notes=hawb_to_bag_for_packing)
            out_path, filename = save_excel_for_part(part_df, itemised_df, idx, invoice_date, session_dir, inv_num, narration_text)
            file_id = f"{session_id}_bag_{idx}"
            download_files[file_id] = {'path': out_path, 'name': filename, 'invoice_number': inv_num}
            file_list.append({'name': filename, 'id': file_id, 'part_number': idx, 'type': 'bag'})

            total_usd = float(part_df['GBP'].sum())
            total_boxes = int(part_df['NO OF PCS'].sum())
            if not itemised_df.empty and 'PCS' in itemised_df.columns:
                pcs_series = pd.to_numeric(itemised_df['PCS'], errors='coerce')
                packing_total_pcs = int(pcs_series.sum()) if not pcs_series.isna().all() else 0
            else:
                packing_total_pcs = 0

            meta = {
                'total_usd': round(total_usd, 2),
                'total_boxes': total_boxes,
                'packing_total_pcs': packing_total_pcs,
                'consignee': heading_consignee,
                'destination': heading_destination
            }
            chamber_path, chamber_filename = generate_chamber_certificate(session_dir, meta, inv_num, invoice_date, idx)
            if chamber_path:
                chamber_file_id = f"{session_id}_bag_chamber_{idx}"
                chamber_download_files[chamber_file_id] = {'path': chamber_path, 'name': chamber_filename, 'invoice_number': inv_num}

        if invoice_numbers:
            save_last_manifest_number(invoice_numbers[-1])

        return jsonify({'files': file_list, 'session_id': session_id})

    else:
        # Only real bag/console-bag labels appear as a note in the packing
        # list - hawb_to_bag already comes from the dedicated field, so
        # every non-empty value here is a genuine marking.
        parts, itemised_parts = split_manifest(merged, bag_map=hawb_to_bag, max_box_per_part=max_box_per_part)
        invoice_date_path = os.path.join(session_dir, 'invoice_date.txt')
        invoice_date = open(invoice_date_path).read().strip() if os.path.exists(invoice_date_path) else datetime.now().strftime('%d-%m-%Y')

        # Bug fix: this branch used to ignore any invoice_numbers sent by the
        # client and always auto-generate ORIG-N, so there was no way to
        # assign a custom invoice number when no bag was assigned. It now
        # honors user-entered invoice numbers the same way the bag-based
        # branch does, falling back to ORIG-N only if none were sent.
        if invoice_numbers is not None and len(invoice_numbers) != len(parts):
            app.logger.error(f"Invoice count mismatch: expected {len(parts)}, got {len(invoice_numbers)}")
            return jsonify({'error': f'Invoice number count mismatch: expected {len(parts)} but received {len(invoice_numbers)}. Please re-enter.'}), 400

        file_list = []
        heading_consignee = "ESHIPPER EXPRESS COURIER"
        heading_destination = "UK"

        for idx, (part_df, itemised_df) in enumerate(zip(parts, itemised_parts), start=1):
            if invoice_numbers and idx - 1 < len(invoice_numbers) and str(invoice_numbers[idx - 1]).strip():
                inv_num = str(invoice_numbers[idx - 1]).strip()
            else:
                inv_num = f"ORIG-{idx}"
            out_path, filename = save_excel_for_part(part_df, itemised_df, idx, invoice_date, session_dir, inv_num)
            file_id = f"{session_id}_orig_{idx}"
            download_files[file_id] = {'path': out_path, 'name': filename, 'invoice_number': inv_num}
            file_list.append({'name': filename, 'id': file_id, 'part_number': idx, 'type': 'orig'})

            total_usd = float(part_df['GBP'].sum())
            total_boxes = int(part_df['NO OF PCS'].sum())
            if not itemised_df.empty and 'PCS' in itemised_df.columns:
                pcs_series = pd.to_numeric(itemised_df['PCS'], errors='coerce')
                packing_total_pcs = int(pcs_series.sum()) if not pcs_series.isna().all() else 0
            else:
                packing_total_pcs = 0

            meta = {
                'total_usd': round(total_usd, 2),
                'total_boxes': total_boxes,
                'packing_total_pcs': packing_total_pcs,
                'consignee': heading_consignee,
                'destination': heading_destination
            }
            chamber_path, chamber_filename = generate_chamber_certificate(session_dir, meta, inv_num, invoice_date, idx)
            if chamber_path:
                chamber_file_id = f"{session_id}_orig_chamber_{idx}"
                chamber_download_files[chamber_file_id] = {'path': chamber_path, 'name': chamber_filename, 'invoice_number': inv_num}

        if invoice_numbers:
            save_last_manifest_number(invoice_numbers[-1])
        elif parts:
            save_last_manifest_number(inv_num)

        return jsonify({'files': file_list, 'session_id': session_id})

@app.route('/download/<file_id>')
@login_required
def download_file(file_id):
    file_info = download_files.get(file_id)
    if not file_info:
        return jsonify({'error': 'File not found'}), 404
    file_path = file_info['path']
    original_name = file_info['name']
    if not os.path.exists(file_path):
        return jsonify({'error': 'File no longer exists'}), 404
    return send_file(file_path, as_attachment=True, download_name=original_name,
                     mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')

@app.route('/chamber_certificate/<file_id>')
@login_required
def chamber_certificate(file_id):
    file_info = chamber_download_files.get(file_id)
    if not file_info:
        return "Certificate not found", 404
    file_path = file_info['path']
    original_name = file_info['name']
    if not os.path.exists(file_path):
        return "File no longer exists", 404
    return send_file(file_path, as_attachment=True, download_name=original_name,
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

@app.route('/download_all/<session_id>')
@login_required
def download_all(session_id):
    session_dir = os.path.join(TEMP_FOLDER, session_id)
    if not os.path.exists(session_dir):
        return "Session not found", 404

    memory_file = io.BytesIO()
    today = datetime.now().strftime('%Y-%m-%d')
    zip_name = f"custom-{today}.zip"

    with zipfile.ZipFile(memory_file, 'w', zipfile.ZIP_DEFLATED) as zf:
        for root, dirs, files in os.walk(session_dir):
            for file in files:
                if (file.startswith('manifest-') or file.startswith('chamber-') or 
                    file.startswith('WT_Box_Report')):
                    file_path = os.path.join(root, file)
                    arcname = file
                    zi = zipfile.ZipInfo(arcname)
                    zi.external_attr = 0o644 << 16
                    with open(file_path, 'rb') as f:
                        zf.writestr(zi, f.read())
    memory_file.seek(0)
    return send_file(memory_file, as_attachment=True, download_name=zip_name,
                     mimetype='application/zip')

@app.route('/manifest.json')
def serve_manifest():
    return send_file('static/manifest.json', mimetype='application/manifest+json')

@app.route('/service-worker.js')
def serve_sw():
    return send_file('static/service-worker.js', mimetype='application/javascript')

@app.route('/offline.html')
def offline_page():
    return render_template('offline.html')

@app.route('/save_wtbox', methods=['POST'])
@login_required
def save_wtbox():
    data = request.get_json()
    session_id = data.get('session_id')
    edited_rows = data.get('data')
    if not session_id or not edited_rows:
        return jsonify({'error': 'Invalid request'}), 400
    df_wt = pd.DataFrame(edited_rows)
    df_wt['Weight'] = pd.to_numeric(df_wt['Weight'], errors='coerce')
    def try_numeric(val):
        if val is None or val == '':
            return None
        try:
            return int(float(val))
        except ValueError:
            try:
                return float(val)
            except ValueError:
                return val
    df_wt['Box'] = df_wt['Box'].apply(try_numeric)
    if 'BagMarking' not in df_wt.columns:
        df_wt['BagMarking'] = ''
    df_wt['BagMarking'] = df_wt['BagMarking'].fillna('')
    total_box = 0
    for val in df_wt['Box']:
        if isinstance(val, (int, float)):
            total_box += val
    total_weight = df_wt['Weight'].sum()
    total_row = pd.DataFrame({
        'HAWB': ['TOTAL'],
        'Details': [''],
        'Box': [total_box],
        'BagMarking': [''],
        'Weight': [round(total_weight, 2)],
        'Country': ['']
    })
    df_wt = pd.concat([df_wt, total_row], ignore_index=True)
    column_order = ['HAWB', 'Details', 'Box', 'BagMarking', 'Weight', 'Country']
    df_wt = df_wt[column_order]
    df_wt = df_wt.rename(columns={'BagMarking': 'Bag Marking'})
    output_filename = f"WT_Box_Report_{session_id}.xlsx"
    output_path = os.path.join(OUTPUT_FOLDER, output_filename)
    with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
        df_wt.to_excel(writer, sheet_name='WT-Box', index=False)
    wb = openpyxl.load_workbook(output_path)
    ws = wb.active
    for column in ws.columns:
        max_len = 0
        col_letter = get_column_letter(column[0].column)
        for cell in column:
            try:
                if cell.value:
                    cell_len = len(str(cell.value))
                    if cell_len > max_len:
                        max_len = cell_len
            except:
                pass
        adjusted = min(max_len + 2, 50)
        ws.column_dimensions[col_letter].width = adjusted
        for cell in column:
            cell.alignment = Alignment(wrap_text=True, vertical='top')
    wb.save(output_path)
    os.chmod(output_path, 0o644)
    session_dir = os.path.join(TEMP_FOLDER, session_id)
    os.makedirs(session_dir, exist_ok=True)
    shutil.copy(output_path, os.path.join(session_dir, f"WT_Box_Report_{session_id}.xlsx"))

    # Persist the Bag Marking assignments (dedicated field, never the Box/pcs
    # column) so generate_manifest / preview_bag_parts can read them even if
    # a later request doesn't resend the full edited row set.
    hawb_to_bag = {}
    for row in edited_rows:
        hawb = row.get('HAWB')
        bag = row.get('BagMarking', row.get('Bag Marking'))
        if bag not in (None, ''):
            hawb_to_bag[hawb] = str(bag).strip()
    save_session_bag_markings(session_dir, hawb_to_bag)

    return send_file(output_path, as_attachment=True, download_name='WT_Box_Report.xlsx')

@app.route('/suggest_bag_markings/<session_id>', methods=['GET'])
@login_required
def suggest_bag_markings_route(session_id):
    """On-demand (re)generation of bag/box marking suggestions, e.g. after
    the user edits weights, or wants to try a different max_weight cap."""
    merged = session_merged_data.get(session_id)
    if merged is None:
        session_dir = os.path.join(TEMP_FOLDER, session_id)
        merged_path = os.path.join(session_dir, 'merged.csv')
        if not os.path.exists(merged_path):
            return jsonify({'error': 'Original data not found'}), 400
        merged = pd.read_csv(merged_path)
    max_weight = request.args.get('max_weight', DEFAULT_MAX_BAG_WEIGHT, type=float)
    suggestions = suggest_bag_markings(merged, max_bag_weight=max_weight)
    return jsonify({'suggestions': suggestions})

# ---------- Last manifest number (continue-from-last-time) ----------
@app.route('/last_manifest_number', methods=['GET'])
@login_required
def last_manifest_number():
    return jsonify({'last_manifest_number': get_last_manifest_number()})

@app.route('/last_manifest_number', methods=['POST'])
@login_required
def update_last_manifest_number():
    data = request.get_json() or {}
    value = str(data.get('manifest_number', '')).strip()
    if not value:
        return jsonify({'error': 'manifest_number is required'}), 400
    save_last_manifest_number(value)
    return jsonify({'success': True, 'last_manifest_number': value})

# ---------- Box-per-part limit (change customs limit without touching code) ----------
@app.route('/last_box_limit', methods=['GET'])
@login_required
def last_box_limit():
    return jsonify({'max_box_per_part': get_last_box_limit()})

@app.route('/last_box_limit', methods=['POST'])
@login_required
def update_last_box_limit():
    data = request.get_json() or {}
    value = data.get('max_box_per_part')
    try:
        value = int(value)
        if value <= 0:
            raise ValueError
    except (TypeError, ValueError):
        return jsonify({'error': 'max_box_per_part must be a positive integer'}), 400
    save_last_box_limit(value)
    return jsonify({'success': True, 'max_box_per_part': value})

# ---------- Register Tracking Blueprint ----------
app.register_blueprint(tracking_bp)

# ---------- Cleanup ----------
def cleanup_temp_files():
    while True:
        time.sleep(3600)
        for folder in os.listdir(TEMP_FOLDER):
            folder_path = os.path.join(TEMP_FOLDER, folder)
            if os.path.isdir(folder_path):
                if time.time() - os.path.getmtime(folder_path) > 3600:
                    shutil.rmtree(folder_path, ignore_errors=True)
if not IS_VERCEL:
    threading.Thread(target=cleanup_temp_files, daemon=True).start()

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)